# -*- coding: utf-8 -*-
"""
Document Assistant — V6 PageIndex
Two-level hierarchical retrieval: page-level summary index + chunk-level store.
Replaces flat SentenceSplitter + BM25 + Reranker pipeline from V5.

CHANGELOG V6
  PI-1  PageIndex: VectorStoreIndex over Page Summary Nodes (Level-1)
  PI-2  chunk_store: deterministic Dict[page_key → List[TextNode]] (Level-2)
  PI-3  Removed BM25Retriever, SentenceTransformerRerank (both variants)
  PI-4  page_retrieve() replaces hybrid_retrieve / hybrid_retrieve_list
  PI-5  fetch_chunks_for_pages() replaces reranker-based node filtering
  PI-6  build_page_context() replaces build_context_with_sources()
  PI-7  Model upgraded: llama-3.1-8b-instant → llama-3.3-70b-versatile
  PI-8  Corporate mode now returns page citations (parity with Academic)
  PI-9  get_similarity_top_k() removed (top_k now inline in ask_question)
  PI-10 Page boundaries: PDF=native fitz, TXT/DOCX=500-word windows
"""

import os
import re
import json
import fitz  # PyMuPDF
from typing import List, Optional, Dict, Tuple, Callable
from dataclasses import dataclass

from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.schema import TextNode
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.postprocessor import SimilarityPostprocessor
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.groq import Groq
from llama_index.readers.file import PyMuPDFReader

import corporate
import academic

# =============================================================================
# GLOBAL SETTINGS
# =============================================================================
Settings.embed_model = HuggingFaceEmbedding(
    model_name="BAAI/bge-small-en-v1.5",
    cache_folder=".embeddings_cache"
)

# =============================================================================
# PROFANITY FILTER  (P1 — unchanged from V5)
# =============================================================================
_BLOCKED_PATTERNS: List[str] = [
    r"f+u+c+k+", r"s+h+i+t+", r"b+i+t+c+h+", r"a+s+s+h+o+l+e+",
    r"c+u+n+t+", r"d+i+c+k+", r"p+i+s+s+", r"c+r+a+p+",
    r"b+a+s+t+a+r+d+", r"n+i+g+g+", r"f+a+g+",
]
_BLOCKED_RE = re.compile(
    "|".join(_BLOCKED_PATTERNS), flags=re.IGNORECASE
)
_SAFE_FALLBACK = (
    "I'm sorry, I can't process that input. "
    "Please rephrase your question professionally."
)


def sanitize_input(text: str) -> Optional[str]:
    """
    Returns None if input is clean.
    Returns a safe generic message if offensive content is detected.
    The original offensive text is NEVER echoed back. (P1)
    """
    if _BLOCKED_RE.search(text):
        return _SAFE_FALLBACK
    return None


# =============================================================================
# RESULT DATACLASS  (unchanged from V5)
# =============================================================================
@dataclass
class AnswerResult:
    """Structured answer with metadata."""
    answer: str
    sources: List[str]


# =============================================================================
# DOCUMENT ASSISTANT
# =============================================================================
class DocumentAssistant:
    """
    PageIndex-based document QA assistant.
    Delegates mode-specific prompt/post-processing to corporate.py / academic.py.
    """

    def __init__(
        self,
        documents: List[Document],
        mode: str = "corporate",
        progress_callback: Optional[Callable[[int, int, str], None]] = None,
    ):
        self.mode      = mode
        self.documents = documents

        # ── PageIndex state ───────────────────────────────────────────────────
        self.page_index:   Optional[VectorStoreIndex] = None
        self.chunk_store:  Dict[str, List[TextNode]]  = {}  # page_key → chunks
        self.page_nodes:   List[TextNode]             = []  # Level-1 nodes

        # ── Retriever (built after index) ─────────────────────────────────────
        self.vector_retriever: Optional[VectorIndexRetriever] = None

        self._build_index(progress_callback)

    # =========================================================================
    # INDEX BUILD  (PI-1, PI-2)
    # =========================================================================

    def _build_index(
        self,
        progress_callback: Optional[Callable] = None
    ) -> None:
        """
        6-step PageIndex build pipeline.
          Step 1 — Extract pages from all documents
          Step 2 — Filter short pages (< 20 words)
          Step 3 — Build Page Summary Nodes (Level-1, first 200 words)
          Step 4 — Build Chunk Nodes (Level-2, SentenceSplitter over full page)
          Step 5 — Build VectorStoreIndex over Level-1 nodes only
          Step 6 — Build chunk_store dict from Level-2 nodes
        """
        total_steps = 6

        # ── Step 1 — Extract pages ────────────────────────────────────────────
        if progress_callback:
            progress_callback(1, total_steps, "Step 1/6  Extracting pages from documents...")

        raw_pages = self._extract_pages(self.documents)

        # ── Step 2 — Filter short pages ───────────────────────────────────────
        if progress_callback:
            progress_callback(2, total_steps, "Step 2/6  Filtering short pages...")

        pages = [
            p for p in raw_pages
            if len(p["text"].split()) >= 20
            and not self._is_structural_page(p["text"])
        ]

        if not pages:
            # Edge case: all pages below threshold (e.g. image-only PDF)
            raise ValueError(
                "No readable text found in the uploaded document(s). "
                "Please check that your files contain selectable text."
            )

        # ── Step 3 — Build Level-1 Page Summary Nodes ─────────────────────────
        if progress_callback:
            progress_callback(3, total_steps, f"Step 3/6  Building page summary nodes ({len(pages)} pages)...")

        self.page_nodes = self._build_page_summary_nodes(pages)

        # ── Step 4 — Build Level-2 Chunk Nodes ───────────────────────────────
        if progress_callback:
            progress_callback(4, total_steps, "Step 4/6  Splitting pages into chunks...")

        self._build_chunk_store(pages)

        # ── Step 5 — Build VectorStoreIndex (Level-1 only) ───────────────────
        if progress_callback:
            progress_callback(5, total_steps, "Step 5/6  Building vector index over page summaries...")

        self.page_index = VectorStoreIndex(self.page_nodes)

        similarity_cutoff = 0.15 if self.mode == "corporate" else 0.18
        self.vector_retriever = VectorIndexRetriever(
            index=self.page_index,
            similarity_top_k=10,          # over-fetch at retrieval; top_k filtered in page_retrieve
            node_postprocessors=[
                SimilarityPostprocessor(similarity_cutoff=similarity_cutoff)
            ]
        )

        # ── Step 6 — chunk_store confirmation ────────────────────────────────
        if progress_callback:
            progress_callback(6, total_steps, f"Step 6/6  Index ready — {len(self.chunk_store)} page buckets built.")

    # =========================================================================
    # PAGE EXTRACTION  (PI-10)
    # =========================================================================

    def _extract_pages(self, documents: List[Document]) -> List[Dict]:
        """
        Dispatch page extraction by filetype.

        Returns list of dicts:
        {
            "text":       str,   # full page text
            "page":       int,   # 1-based page number
            "filename":   str,
            "filetype":   str,   # "pdf" | "txt" | "docx"
            "source":     str,   # full file path
        }
        """
        all_pages = []

        # Group documents by source file so we can process PDF files via fitz
        # using the source path rather than reconstructed text.
        seen_sources = {}
        for doc in documents:
            src = doc.metadata.get("source", "")
            if src not in seen_sources:
                seen_sources[src] = doc  # first doc per source carries metadata

        for source_path, doc in seen_sources.items():
            filetype = doc.metadata.get("filetype", "").lower()
            filename = doc.metadata.get("filename", os.path.basename(source_path))

            if filetype == "pdf":
                pages = self._extract_pages_pdf(source_path, filename)
            elif filetype == "txt":
                # Collect full text from all Document objects for this source
                full_text = " ".join(
                    d.text for d in documents
                    if d.metadata.get("source", "") == source_path
                )
                pages = self._extract_pages_txt(full_text, filename, source_path)
            elif filetype == "docx":
                full_text = " ".join(
                    d.text for d in documents
                    if d.metadata.get("source", "") == source_path
                )
                pages = self._extract_pages_docx(full_text, filename, source_path)
            else:
                # Unknown filetype — treat as plain text
                full_text = " ".join(
                    d.text for d in documents
                    if d.metadata.get("source", "") == source_path
                )
                pages = self._extract_pages_txt(full_text, filename, source_path)

            all_pages.extend(pages)

        return all_pages

    def _extract_pages_pdf(
        self, source_path: str, filename: str
    ) -> List[Dict]:
        """Extract pages from PDF using fitz (native page boundaries)."""
        pages = []
        try:
            doc = fitz.open(source_path)
            for i, page in enumerate(doc):
                text = page.get_text("text").strip()
                pages.append({
                    "text":     text,
                    "page":     i + 1,          # 1-based
                    "filename": filename,
                    "filetype": "pdf",
                    "source":   source_path,
                })
            doc.close()
        except Exception as e:
            print(f"PDF extraction error ({filename}): {str(e)}")
        return pages

    def _extract_pages_txt(
        self, text: str, filename: str, source_path: str
    ) -> List[Dict]:
        """Split plain text into 500-word page windows."""
        page_texts = self._split_into_pages(text, words_per_page=500)
        return [
            {
                "text":     page_text,
                "page":     i + 1,
                "filename": filename,
                "filetype": "txt",
                "source":   source_path,
            }
            for i, page_text in enumerate(page_texts)
        ]

    def _extract_pages_docx(
        self, text: str, filename: str, source_path: str
    ) -> List[Dict]:
        """Split DOCX text into 500-word page windows."""
        page_texts = self._split_into_pages(text, words_per_page=500)
        return [
            {
                "text":     page_text,
                "page":     i + 1,
                "filename": filename,
                "filetype": "docx",
                "source":   source_path,
            }
            for i, page_text in enumerate(page_texts)
        ]

    @staticmethod
    def _split_into_pages(text: str, words_per_page: int = 500) -> List[str]:
        """
        Split text into fixed-size word-count windows.
        No overlap at page level — overlap lives at Level-2 chunk level.
        """
        words = text.split()
        if not words:
            return []
        return [
            " ".join(words[i: i + words_per_page])
            for i in range(0, len(words), words_per_page)
        ]

    # =========================================================================
    # NODE BUILDERS  (PI-1, PI-2)
    # =========================================================================

    def _build_page_summary_nodes(self, pages: List[Dict]) -> List[TextNode]:
        """
        Build Level-1 Page Summary Nodes.
        Text = first 200 words of the page (topic proxy for embedding).
        Full page text lives in Level-2 chunk_store only.
        """
        nodes = []
        for p in pages:
            summary_text = " ".join(p["text"].split()[:400])
            page_key     = f"{p['filename']}::page::{p['page']}"

            node = TextNode(
                text=summary_text,
                metadata={
                    "page":            p["page"],
                    "filename":        p["filename"],
                    "filetype":        p["filetype"],
                    "source":          p["source"],
                    "level":           "page",
                    "page_key":        page_key,
                    "page_word_count": len(p["text"].split()),
                }
            )
            nodes.append(node)
        return nodes

    def _build_chunk_store(self, pages: List[Dict]) -> None:
        """
        Build Level-2 chunk_store.
        SentenceSplitter applied over full page text.
        Chunks stored by page_key — never indexed, fetched deterministically.
        """
        chunk_size    = 600 if self.mode == "corporate" else 500
        chunk_overlap = 150
        splitter      = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )

        for p in pages:
            page_key = f"{p['filename']}::page::{p['page']}"
            # Create a temporary Document so SentenceSplitter can parse it
            tmp_doc  = Document(
                text=p["text"],
                metadata={
                    "page":     p["page"],
                    "filename": p["filename"],
                    "filetype": p["filetype"],
                    "source":   p["source"],
                    "level":    "chunk",
                    "page_key": page_key,
                }
            )
            chunk_nodes = splitter.get_nodes_from_documents([tmp_doc])

            # Apply smart truncation (inherited from V5)
            max_chars = 2000 if self.mode == "corporate" else 1500
            chunk_nodes = self._smart_truncate_nodes(chunk_nodes, max_chars)

            self.chunk_store[page_key] = chunk_nodes
            
    
    def _is_structural_page(self, page_text: str) -> bool:
        """
        Returns True if page is structural — TOC, index, ad, or header-only page.
        Structural pages have majority of lines being very short (5 words or less).
        Filtered out during index build to prevent TOC content polluting retrieval.
        """
        lines = [l.strip() for l in page_text.splitlines() if l.strip()]
        if not lines:
            return True
        short_lines = sum(1 for l in lines if len(l.split()) <= 5)
        return (short_lines / len(lines)) >= 0.6

    
    # =========================================================================
    # SMART TRUNCATE  (unchanged from V5)
    # =========================================================================

    @staticmethod
    def _smart_truncate_nodes(nodes: List, max_chars: int) -> List:
        """Truncate nodes at sentence boundaries."""
        processed = []
        for node in nodes:
            if len(node.text) <= max_chars:
                processed.append(node)
                continue
            sentences = re.split(r'(?<=[.!?])\s+', node.text)
            truncated = ""
            for sent in sentences:
                if len(truncated) + len(sent) + 1 <= max_chars:
                    truncated += sent + " "
                else:
                    break
            if len(truncated.strip()) >= 50:
                node.text = truncated.strip()
            processed.append(node)
        return processed

    # =========================================================================
    # RETRIEVAL  (PI-4, PI-5)
    # =========================================================================

    def page_retrieve(
        self,
        queries: List[str],
        top_k: int
    ) -> List[TextNode]:
        """
        Level-1 retrieval — query the Page Summary index.
        Deduplicates by page_key, keeps highest score per page.
        Returns up to top_k page nodes sorted by score descending.
        """
        best: Dict[str, Tuple[float, TextNode]] = {}
        # score → node

        for query in queries:
            try:
                results = self.vector_retriever.retrieve(query)
            except Exception as e:
                print(f"Retrieval error for query '{query}': {str(e)}")
                continue

            for node_with_score in results:
                node  = node_with_score.node
                score = node_with_score.score or 0.0
                key   = node.metadata.get("page_key", "")

                if not key:
                    # Fallback key if metadata missing
                    key = (
                        f"{node.metadata.get('filename','unknown')}"
                        f"::page::{node.metadata.get('page', 0)}"
                    )

                if key not in best or score > best[key][0]:
                    best[key] = (score, node)

        # Sort by score descending, take top_k
        sorted_nodes = sorted(best.values(), key=lambda x: x[0], reverse=True)
        return [node for _, node in sorted_nodes[:top_k]]

    def fetch_chunks_for_pages(
        self,
        page_nodes: List[TextNode]
    ) -> List[TextNode]:
        """
        Level-2 fetch — retrieve all chunks for matched pages from chunk_store.
        Returns chunks ordered by page number (ascending) for coherent context.
        """
        all_chunks = []

        # Sort page_nodes by page number before fetching
        sorted_pages = sorted(
            page_nodes,
            key=lambda n: n.metadata.get("page", 9999)
        )

        for node in sorted_pages:
            page_key = node.metadata.get("page_key", "")
            if not page_key:
                page_key = (
                    f"{node.metadata.get('filename','unknown')}"
                    f"::page::{node.metadata.get('page', 0)}"
                )
            chunks = self.chunk_store.get(page_key, [])
            all_chunks.extend(chunks)

        return all_chunks

    # =========================================================================
    # CONTEXT ASSEMBLY  (PI-6)
    # =========================================================================

    def build_page_context(
        self,
        chunks: List[TextNode],
        question: str,
        wide: bool = False
    ) -> Tuple[str, Dict[str, List[str]]]:
        """
        Assemble context string and source map from Level-2 chunk nodes.
        Academic mode: page-tagged context for accurate LLM self-reporting.
        Corporate mode: plain context (no tags), sourcemap for coverage indicator.
        """
        # ── Character budget (inherited from V5) ─────────────────────────────
        question_words = len(question.split())
        if question_words <= 5:
            base_chars = 4000
        elif question_words <= 12:
            base_chars = 7000
        else:
            base_chars = 9000

        if wide:
            base_chars = max(base_chars, 20000)

        max_context_chars = int(base_chars * 1.5) if self.mode == "academic" else base_chars

        # ── Assemble ──────────────────────────────────────────────────────────
        context_parts: List[str] = []
        source_map:    Dict[str, List[str]] = {}
        total_chars    = 0

        for node in chunks:
            txt = node.text.strip()
            if not txt:
                continue
            if total_chars + len(txt) > max_context_chars:
                break

            meta     = node.metadata or {}
            filename = (
                meta.get("filename")
                or meta.get("file_name")
                or meta.get("file_path", "")
            )
            filename = os.path.basename(filename) if filename else "document"
            page     = meta.get("page", "")

            # ── Context string format ─────────────────────────────────────────
            if self.mode == "academic":
                # Page-tagged so LLM can accurately self-report pages_used
                header = f"[Page {page} — {filename}]" if page else f"[{filename}]"
                context_parts.append(f"{header}\n{txt}")
            else:
                # Corporate: light page tagging to prevent entity blending
                if page:
                    context_parts.append(f"[Page {page} — {filename}]\n{txt}")
                else:
                    context_parts.append(txt)


            # ── Source map (both modes) ───────────────────────────────────────
            source_str = filename
            if page and str(page).strip().isdigit():
                source_str = f"{filename} p.{page}"
            source_map.setdefault(source_str, []).append(txt)

            total_chars += len(txt)

        context = "\n\n---\n\n".join(context_parts)
        return context, source_map

    # =========================================================================
    # SUMMARY GENERATION  (unchanged from V5 except model string)
    # =========================================================================

    def generate_summary(
        self,
        groq_api_key: str,
        progress_callback: Optional[Callable[[int, int, str], None]] = None
    ) -> str:
        """Generate a concise document summary."""
        try:
            total_steps = 3
            if progress_callback:
                progress_callback(1, total_steps, "Initializing language model...")

            llm = Groq(
                model="llama-3.1-8b-instant",      # PI-7
                api_key=groq_api_key,
                temperature=0.2,
                max_tokens=400
            )

            if progress_callback:
                progress_callback(2, total_steps, "Analyzing document content...")

            filtered_docs = self._filter_summary_docs()
            if not filtered_docs:
                return "Summary unavailable — document contains primarily structural content."

            context = self._build_summary_context(filtered_docs, max_chars=8000)

            if progress_callback:
                progress_callback(3, total_steps, "Generating summary...")

            prompt   = self._get_summary_prompt(context)
            response = llm.complete(prompt)
            summary  = str(response).strip()

            if len(summary.split()) < 30:
                return "Summary unavailable — insufficient substantive content."

            return summary

        except Exception as e:
            return f"Summary generation failed: {str(e)}"

    def _filter_summary_docs(self) -> List[Document]:
        """Filter out structural/boilerplate documents before summarisation."""
        if self.mode == "corporate":
            skip_patterns = ["table of contents", "objectives", "copyright"]
        else:
            skip_patterns = [
                "learning outcomes", "unit objectives", "check your progress",
                "terminal questions", "glossary", "key words"
            ]
        filtered = []
        for doc in self.documents:
            text_lower  = doc.text.lower()
            junk_matches = sum(1 for p in skip_patterns if p in text_lower)
            if junk_matches < 2:
                filtered.append(doc)
        return filtered

    def _build_summary_context(
        self, docs: List[Document], max_chars: int
    ) -> str:
        """First 1000 chars of each doc, up to max_chars total."""
        context = ""
        for doc in docs[:10]:
            chunk = doc.text[:1000].strip()
            if len(context) + len(chunk) > max_chars:
                break
            context += chunk
        return context.strip()

    def _get_summary_prompt(self, context: str) -> str:
        base = (
            "Create a single-paragraph summary (100–120 words max). "
            "Focus ONLY on the main topic and key themes. "
            "Use natural, flowing prose — no lists or segmentation. "
            "Do NOT start with 'This document discusses' or 'The document'.\n\n"
        )
        specific = (
            "Focus: What the document covers and why it matters."
            if self.mode == "corporate"
            else "Focus: Core arguments and contributions."
        )
        return f"{base}{specific}\n\nCONTEXT:\n{context}"

    # =========================================================================
    # QUESTION ANSWERING
    # =========================================================================

    def ask_question(
        self,
        question: str,
        groq_api_key: str,
        return_metadata: bool = False
    ) -> AnswerResult:
        """Answer a question grounded in the loaded documents."""

        # ── P1 Profanity filter ───────────────────────────────────────────────
        blocked = sanitize_input(question)
        if blocked:
            result = AnswerResult(answer=blocked, sources=[])
            return result if return_metadata else result.answer

        # ── Basic validation ──────────────────────────────────────────────────
        if not question or len(question.strip()) < 3:
            error = "Please ask a more specific question."
            return AnswerResult(error, []) if return_metadata else error

        if not self.page_index:
            error = "Error: Index not initialized."
            return AnswerResult(error, []) if return_metadata else error

        try:
            # ── Normalize + classify ──────────────────────────────────────────
            normalized_question = self.normalize_question(question)

            # Append grounding anchor for attribution-scoped list questions only
            _ATTRIBUTION_TRIGGERS = {
                "name a few", "name some", "list a few", "list some",
                "give a few", "give some", "mention a few", "mention some"
            }   


            q_lower = normalized_question.lower()
            if any(t in q_lower for t in _ATTRIBUTION_TRIGGERS):
                if "mentioned in this doc" not in q_lower:
                    normalized_question = normalized_question + " mentioned in this document"

            is_list_q           = self.is_list_question(normalized_question)
            is_simple           = self.is_simple_question(normalized_question)
            is_comparison       = normalized_question.strip().lower().startswith("compare")
            expanded_queries    = self.expand_query(normalized_question)

            # ── PI-4 PageIndex retrieval ──────────────────────────────────────
            if is_list_q:
                top_k_pages = 8
            elif is_comparison:
                top_k_pages = 5
            else:
                top_k_pages = 3

            page_nodes     = self.page_retrieve(expanded_queries, top_k=top_k_pages)
            retrieved_chunks = self.fetch_chunks_for_pages(page_nodes)

            if not retrieved_chunks:
                result = AnswerResult(
                    answer="This information is not covered in the provided documents.",
                    sources=[]
                )
                return result if return_metadata else result.answer

            # ── PI-6 Context assembly ─────────────────────────────────────────
            wide = is_list_q or is_comparison
            context, source_map = self.build_page_context(
                retrieved_chunks, normalized_question, wide=wide
            )

            # ── LLM setup ────────────────────────────────────────────────────
            if self.mode == "corporate":
                max_tokens = corporate.get_token_budget(normalized_question)
            else:
                max_tokens = academic.get_token_budget(normalized_question)

            llm = Groq(
                model="llama-3.1-8b-instant",      # PI-7
                api_key=groq_api_key,
                temperature=0.0,
                max_tokens=max_tokens
            )

            # ── Prompt build ──────────────────────────────────────────────────
            if self.mode == "corporate":
                prompt = corporate.get_answer_prompt(
                    context,
                    normalized_question,
                    is_list_q=is_list_q,
                    is_comparison=is_comparison,
                    is_simple=is_simple,
                )
            else:
                prompt = academic.get_answer_prompt(context, normalized_question)


            # ── LLM call ─────────────────────────────────────────────────────
            response   = llm.complete(prompt)
            raw_output = str(response).strip()

            # ── Parse response ────────────────────────────────────────────────
            pages_used = []

            if self.mode in ("academic", "corporate"):
                try:
                    # Strip markdown fences and any text before the first {
                    clean_json = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw_output, flags=re.DOTALL).strip()
                    # Find the first { and last } to extract pure JSON
                    start = clean_json.find("{")
                    end   = clean_json.rfind("}") + 1
                    if start != -1 and end > start:
                        clean_json = clean_json[start:end]
                    parsed     = json.loads(clean_json)
                    answer     = parsed.get("answer", "").strip()
                    pages_used = [
                        int(p) for p in parsed.get("pages_used", [])
                    ]
                except (json.JSONDecodeError, KeyError, ValueError):
                    # Graceful fallback — treat entire output as plain answer
                    answer     = raw_output
                    pages_used = []


            if not answer or len(answer.strip()) < 10:
                error = "Unable to generate answer. Please rephrase."
                return AnswerResult(error, []) if return_metadata else error

            # ── Post-processing ───────────────────────────────────────────────
            answer     = self._post_process_answer(answer, normalized_question)
            is_negative = self._is_negative_response(answer)

            # ── Source assignment ─────────────────────────────────────────────
            if is_negative:
                sources = []                   
            elif pages_used:
                # Both modes: use LLM self-reported pages
                sources = [f"Page {p}" for p in sorted(set(pages_used))]
            else:
                # Fallback if JSON parse failed
                sources = list(source_map.keys())[:3]

            result = AnswerResult(answer=answer, sources=sources)
            return result if return_metadata else result.answer

        except Exception:
            import traceback
            traceback.print_exc()
            error = "Error processing question. Please try rephrasing."
            return AnswerResult(error, []) if return_metadata else error

    # =========================================================================
    # QUESTION CLASSIFICATION  (unchanged from V5)
    # =========================================================================

    def is_list_question(self, question: str) -> bool:
        q = question.lower()
        comparison_phrases = [
            "difference between", "differences between", "distinguish between",
            "compare", "comparison between", "contrast between", "vs", "versus"
        ]
        if any(p in q for p in comparison_phrases):
            return False
        indicators = [
            "what are", "list", "types of", "kinds of", "styles of",
            "phases", "stages", "steps", "eras", "periods", "categories",
            "enumerate", "name all", "name the", "name", "give",
            "mention", "outline", "what are the"
        ]
        return any(i in q for i in indicators)

    def is_simple_question(self, question: str) -> bool:
        """Detect simple single-fact questions."""
        q = question.strip().lower()
        if q.startswith("compare") or "vs" in q or "versus" in q:
            return False
        if self.is_list_question(question):
            return False
        simple_patterns = [
            r"what is", r"who is", r"when (was|is|did)",
            r"where (is|was)", r"how (many|much|old|long|tall)",
        ]
        return any(re.search(p, q) for p in simple_patterns)

    def normalize_question(self, question: str) -> str:
        """
        Rewrite surface variants into canonical forms.
        Comparison rewrites → 'compare X and Y'
        """
        q = question.strip().lower()

        # Pattern 1: what are the differences between X and Y
        m = re.search(
            r"what\s+are?\s+the\s+differences?\s+between\s+(.+?)\s+and\s+(.+)",
            q
        )
        if m:
            return f"compare {m.group(1).strip()} and {m.group(2).strip()}"

        # Pattern 2: how does X differ from Y
        m = re.search(r"how\s+does\s+(.+?)\s+differ\s+from\s+(.+)", q)
        if m:
            return f"compare {m.group(1).strip()} and {m.group(2).strip()}"

        # Pattern 3: difference between X and Y (no leading what/how)
        m = re.search(r"differences?\s+between\s+(.+?)\s+and\s+(.+)", q)
        if m:
            return f"compare {m.group(1).strip()} and {m.group(2).strip()}"

        return question

    def expand_query(self, question: str) -> List[str]:
        """
        Expand a question into multiple search queries
        to improve Level-1 page retrieval coverage.
        """
        queries = [question]
        q       = question.lower().strip()

        if q.startswith("compare"):
            rest = question[8:].strip()
            if " and " in rest:
                parts     = rest.split(" and ", 1)
                subject_a = parts[0].strip()
                subject_b = parts[1].strip()
                queries   = [
                    f"what is {subject_a}",
                    f"what is {subject_b}",
                    f"{subject_a} characteristics",
                    f"{subject_b} characteristics",
                ]

        elif q.startswith("who is") or q.startswith("who was") or q.startswith("who"):
            words = question.split()
            verb_triggers = {"sang", "played", "wrote", "written", "directed",
                             "narrated", "performed", "recorded", "composed",
                             "created", "founded", "invented", "discovered"}
            if len(words) > 1 and words[1].lower() in verb_triggers:
                # Attribution question — who sang/wrote/founded X
                work = " ".join(words[2:]).strip()
                queries = [
                    work,
                    f"{work} and what is it about",
                    f"about {work}",
                    f"{words[1]} {work}",
                    f"{work} background",
                ]
            else:
                # Identity question — who is/was X
                term = " ".join(words[1:]).strip()
                queries = [
                    f"called {term}",
                    f"known as {term}",
                    f"referred to as {term}",
                    f"{term} history background",
                    term,
                ]


        elif q.startswith("explain") or q.startswith("expand on"):
            term    = " ".join(q.split()[1:]).strip()
            queries = [
                f"how did {term}",
                f"impact of {term}",
                f"role of {term}",
                f"{term} influence",
                term,
            ]

        elif q.startswith("what is"):
            term    = question[8:].strip()
            queries = [
                f"define {term}",
                f"explain {term}",
                f"{term} meaning",
                f"{term} characteristics",
                f"{term} description",
                f"the {term}",
            ]

        elif q.startswith("what are"):
            term    = question[9:].strip()
            queries = [
                f"list of {term}",
                f"types of {term}",
            ]

        elif q.startswith("how to"):
            queries = [
                question.replace("how to", "steps for"),
                question.replace("how to", "process of"),
            ]
        elif q.startswith("how did") or q.startswith("how does") or q.startswith("how do"):
            term = " ".join(q.split()[2:]).strip()
            queries = [
                f"impact of {term}",
                f"influence of {term}",
                f"role of {term}",
                term,
            ]

        elif q.startswith("why"):
            queries.append(question.replace("why", "reasons for"))

        # Always append keyword-only fallback query
        filler    = {"the", "a", "an", "is", "are", "was", "were"}
        key_terms = " ".join(w for w in q.split() if w not in filler)
        if key_terms != q:
            queries.append(key_terms)

        return queries[:6]


    # =========================================================================
    # POST-PROCESSING  (unchanged from V5)
    # =========================================================================

    def _post_process_answer(self, answer: str, question: str) -> str:
        """
        Full post-processing pipeline:
          Step 1 — P4: Strip internal section references
          Step 2 — P5: Deduplicate near-identical sentences
          Step 3 — Mode-specific length/brevity rules
        """
        # Step 1
        answer = self._strip_internal_refs(answer)
        answer = answer[0].upper() + answer[1:] if answer else answer
        
        # Step 2
        answer = self._deduplicate_sentences(answer)

        # Step 3
        is_simple = self.is_simple_question(question)
        is_list_q = self.is_list_question(question)

        if is_simple and not is_list_q:
            sentences = re.split(r'(?<=[.!?])\s+', answer)
            q_lower   = question.lower().strip()
            if any(t in q_lower for t in ["name", "list"]):
                answer = sentences[0].strip() if sentences else answer
            else:
                answer = " ".join(sentences[:2]).strip() if len(sentences) >= 2 else answer
            if not answer.endswith((".", "!", "?")):
                answer += "."

        if self.mode == "corporate":
            answer = corporate.post_process(answer, question)
        else:
            answer = academic.post_process(answer, question)

        return answer.strip()

    def _is_negative_response(self, answer: str) -> bool:
        """Detect responses where the LLM signals it couldn't find information."""
        negative_phrases = [
            "not covered", "not mentioned", "no information",
            "cannot find", "not available", "not provided",
            "does not contain", "i don't know", "i do not know",
            "not found in", "outside the scope",
        ]
        lowered = answer.strip().lower()
        return any(phrase in lowered for phrase in negative_phrases)

    @staticmethod
    def _strip_internal_refs(text: str) -> str:
        """
        P4 — Remove internal document references that leak into answers.
        Covers: Section A.2, Chapter 4, Unit 2, Appendix B,
                Page 12, p.12, Figure 3.1, Table 2.3, see above/below/page N.
        """
        patterns = [
            r"\b[Ss]ection\s+[A-Z0-9]+(?:\.[A-Z0-9]+)*",
            r"\b[Cc]hapter\s+\d+",
            r"\b[Uu]nit\s+\d+",
            r"\b[Aa]ppendix\s+[A-Z]",
            r"\b[Pp]age\s+\d+",
            r"\bp\.\s*\d+",
            r"\b[Ff]igure\s+\d+(?:\.\d+)?",
            r"\b[Tt]able\s+\d+(?:\.\d+)?",
            r"\b(?:see\s+)?(?:above|below|page)\s+\d*",
        ]
        cleaned = text
        for pat in patterns:
            cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE)

        # Collapse multiple spaces / orphaned commas
        cleaned = re.sub(r"\s{2,}", " ", cleaned)
        cleaned = re.sub(r",\s*,", ",", cleaned)
        cleaned = re.sub(r",\s*\.", ".", cleaned)
        return cleaned.strip()

    @staticmethod
    def _deduplicate_sentences(text: str) -> str:
        """
        P5 — Remove near-duplicate sentences to reduce verbosity/repetition.
        Handles both prose and numbered lists.
        """
        lines = text.splitlines()

        segments: List[str] = []
        prose_buffer = ""

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if prose_buffer.strip():
                    segments.extend(re.split(r'(?<=[.!?])\s+', prose_buffer.strip()))
                    prose_buffer = ""
                continue
            # Numbered list items stay whole
            if re.match(r"^\d+[\.\)]\s+\S", stripped):
                if prose_buffer.strip():
                    segments.extend(re.split(r'(?<=[.!?])\s+', prose_buffer.strip()))
                    prose_buffer = ""
                segments.append(stripped)
            else:
                prose_buffer += " " + line

        if prose_buffer.strip():
            segments.extend(re.split(r'(?<=[.!?])\s+', prose_buffer.strip()))

        def normalise(s: str) -> str:
            return re.sub(r'[^a-z0-9]', '', s.lower().strip())

        seen_norms:     List[str] = []
        unique_segments: List[str] = []

        for seg in segments:
            seg  = seg.strip()
            if not seg:
                continue
            norm = normalise(seg)

            # Skip if this is a substring of something already kept
            is_subset = any(
                norm in existing and len(norm) > 20
                for existing in seen_norms
            )
            if is_subset:
                continue

            # Replace an existing shorter sentence with this longer one
            replaced = False
            for i, existing in enumerate(seen_norms):
                if existing in norm and len(existing) > 20:
                    seen_norms[i]      = norm
                    unique_segments[i] = seg
                    replaced           = True
                    break

            if not replaced:
                seen_norms.append(norm)
                unique_segments.append(seg)

        # Re-join — list items get their own line, prose gets a space
        result_parts: List[str] = []
        for seg in unique_segments:
            if re.match(r"^\d+[\.\)]\s", seg):
                result_parts.append(seg)
            else:
                result_parts.append(seg)

        return "\n".join(result_parts).strip()


# =============================================================================
# DOCUMENT LOADER  (unchanged from V5)
# =============================================================================

def load_documents(
    file_paths: List[str],
    progress_callback: Optional[Callable[[int, int, str], None]] = None
) -> List[Document]:
    """
    Load PDF, TXT, or DOCX files into LlamaIndex Document objects
    with rich metadata: filename, page, source, filetype.
    """
    documents: List[Document] = []
    total_files = len(file_paths)

    for idx, path in enumerate(file_paths, 1):
        if progress_callback:
            progress_callback(idx, total_files, f"Loading {os.path.basename(path)}...")

        filename = os.path.basename(path)

        try:
            if path.lower().endswith(".pdf"):
                reader = PyMuPDFReader()
                docs   = reader.load_data(file_path=path)

                for i, doc in enumerate(docs):
                    doc.metadata.update({
                        "filename": filename,
                        "page":     i + 1,
                        "source":   path,
                        "filetype": "pdf",
                    })
                    documents.append(doc)

            elif path.lower().endswith(".txt"):
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                documents.append(Document(
                    text=text,
                    metadata={
                        "filename": filename,
                        "source":   path,
                        "filetype": "txt",
                    }
                ))

            elif path.lower().endswith(".docx"):
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(path)
                    text     = "\n".join(
                        para.text for para in docx_doc.paragraphs
                        if para.text.strip()
                    )
                except ImportError:
                    # Fallback: raw binary read
                    with open(path, "rb") as f:
                        raw  = f.read()
                    text = raw.decode("utf-8", errors="ignore")

                documents.append(Document(
                    text=text,
                    metadata={
                        "filename": filename,
                        "source":   path,
                        "filetype": "docx",
                    }
                ))

        except Exception as e:
            print(f"Error loading {filename}: {str(e)}")
            continue

    return documents